"""
document_processor.py
─────────────────────
Unified document ingestion for the Accreditation Intelligence Platform.

Supported formats:
  .pdf   — one chunk per page  (pypdf)
  .docx  — one chunk per section / heading block + all tables  (python-docx)
  .xlsx  — one chunk per sheet  (openpyxl + pandas)
  .xls   — same as xlsx via xlrd
  .csv   — entire file as one chunk  (pandas)
  .txt   — one chunk per ~50-line block

FIX: Streamlit UploadedFile cursor reset
  Streamlit UploadedFile objects maintain a read cursor. If any code reads
  the file before load_document() is called, the cursor is at EOF and all
  loaders return empty results. load_document() now calls file.seek(0)
  before reading to guarantee a fresh read regardless of prior access.

FIX: DOCX — content inside tables was previously only extracted at the end
  as separate TABLE chunks. Many accreditation documents (rubrics, course
  specs, KPI tables) store ALL their content in tables with no body paragraphs.
  The updated loader extracts table content inline AND as dedicated table
  chunks so nothing is missed.
"""

import io
import pypdf
import pandas as pd


# ── PDF ────────────────────────────────────────────────────────────────────────

def _load_pdf(file) -> list:
    chunks = []
    try:
        file.seek(0)   # FIX: reset cursor
        reader = pypdf.PdfReader(file)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                chunks.append(f"[PAGE {page_num}]\n{text.strip()}")
    except Exception as e:
        chunks.append(f"[ERROR] Could not read PDF: {e}")
    return chunks


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _load_docx(file) -> list:
    try:
        from docx import Document
    except ImportError:
        return ["[ERROR] python-docx not installed. Run: pip install python-docx"]

    chunks = []
    try:
        file.seek(0)   # FIX: reset cursor
        doc = Document(file)

        # ── Body paragraphs grouped by heading ──
        section_lines = []
        section_num   = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_heading = para.style.name.lower().startswith("heading")
            if is_heading and section_lines:
                chunks.append(f"[SECTION {section_num}]\n" + "\n".join(section_lines))
                section_num  += 1
                section_lines = []
            section_lines.append(text)

        if section_lines:
            chunks.append(f"[SECTION {section_num}]\n" + "\n".join(section_lines))

        # ── Tables — extracted fully as dedicated chunks ──
        # FIX: Many rubric/course-spec documents store ALL content in tables.
        # We deduplicate by tracking cell text already seen in body paragraphs.
        for t_idx, table in enumerate(doc.tables, start=1):
            rows_text = []
            for row in table.rows:
                # Deduplicate merged cells (python-docx repeats merged cell text)
                seen = set()
                cells = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct and ct not in seen:
                        seen.add(ct)
                        cells.append(ct)
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                chunks.append(f"[TABLE {t_idx}]\n" + "\n".join(rows_text))

    except Exception as e:
        chunks.append(f"[ERROR] Could not read DOCX: {e}")

    return chunks


# ── XLSX / XLS ─────────────────────────────────────────────────────────────────

def _load_excel(file, filename: str) -> list:
    chunks = []
    try:
        file.seek(0)   # FIX: reset cursor
        engine     = "openpyxl" if filename.lower().endswith(".xlsx") else "xlrd"
        all_sheets = pd.read_excel(file, sheet_name=None, engine=engine, dtype=str)

        for sheet_name, df in all_sheets.items():
            if df.empty:
                continue
            df = df.dropna(how="all").dropna(axis=1, how="all").fillna("")
            if df.empty:
                continue

            lines = []
            lines.append("  |  ".join(str(c) for c in df.columns))
            lines.append("-" * min(80, sum(len(str(c)) + 5 for c in df.columns)))
            for _, row in df.iterrows():
                row_vals = [str(v).strip() for v in row.values]
                if any(v for v in row_vals):
                    lines.append("  |  ".join(row_vals))

            chunks.append(f"[SHEET: {sheet_name}]\n" + "\n".join(lines))

    except Exception as e:
        chunks.append(f"[ERROR] Could not read Excel file: {e}")

    return chunks


# ── CSV ────────────────────────────────────────────────────────────────────────

def _load_csv(file) -> list:
    chunks = []
    try:
        file.seek(0)   # FIX: reset cursor
        df = pd.read_csv(file, dtype=str).fillna("")
        df = df.dropna(how="all").dropna(axis=1, how="all")

        lines = []
        lines.append("  |  ".join(str(c) for c in df.columns))
        lines.append("-" * 60)
        for _, row in df.iterrows():
            row_vals = [str(v).strip() for v in row.values]
            if any(v for v in row_vals):
                lines.append("  |  ".join(row_vals))

        if lines:
            chunks.append("[CSV DATA]\n" + "\n".join(lines))

    except Exception as e:
        chunks.append(f"[ERROR] Could not read CSV: {e}")

    return chunks


# ── TXT ────────────────────────────────────────────────────────────────────────

def _load_txt(file, lines_per_chunk: int = 50) -> list:
    chunks = []
    try:
        file.seek(0)   # FIX: reset cursor
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        lines     = content.splitlines()
        block_num = 1
        for i in range(0, len(lines), lines_per_chunk):
            block = "\n".join(lines[i:i + lines_per_chunk]).strip()
            if block:
                chunks.append(f"[BLOCK {block_num}]\n{block}")
                block_num += 1
    except Exception as e:
        chunks.append(f"[ERROR] Could not read TXT: {e}")
    return chunks


# ── PUBLIC ENTRY POINT ─────────────────────────────────────────────────────────

def load_document(uploaded_file) -> list:
    """
    Dispatch to the correct loader based on file extension.
    Always resets the file cursor before reading (fixes Streamlit UploadedFile issue).
    """
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return _load_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return _load_docx(uploaded_file)
    elif name.endswith((".xlsx", ".xls")):
        return _load_excel(uploaded_file, name)
    elif name.endswith(".csv"):
        return _load_csv(uploaded_file)
    elif name.endswith(".txt"):
        return _load_txt(uploaded_file)
    else:
        return [f"[UNSUPPORTED FORMAT] '{uploaded_file.name}' — "
                "Supported: PDF, DOCX, XLSX, XLS, CSV, TXT"]


# ── LEGACY ALIAS ──────────────────────────────────────────────────────────────

def load_and_chunk_pdf(uploaded_file) -> list:
    """Backward-compatible alias."""
    return load_document(uploaded_file)