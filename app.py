import streamlit as st
import os
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from anthropic import Anthropic, NotFoundError
from pypdf import PdfReader

# --- CONFIGURATION ---
st.set_page_config(page_title="NCAAA Dentistry Accreditation AI", page_icon="🦷", layout="wide")

# --- SECURE KEY HANDLING ---
try:
    if "ANTHROPIC_API_KEY" in st.secrets:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    else:
        load_dotenv()
        api_key = os.getenv("ANTHROPIC_API_KEY")
except FileNotFoundError:
    load_dotenv()
    api_key = os.getenv("ANTHROPIC_API_KEY")

if not api_key:
    st.error("🚨 Configuration Error: ANTHROPIC_API_KEY is missing. Please add it to Streamlit Secrets.")
    st.stop()

client = Anthropic(api_key=api_key)

# --- MODELS ---
# Tries the smartest model first. Falls back to fast model if account is restricted.
SMART_MODEL = "claude-3-5-sonnet-20240620"
FAST_MODEL = "claude-3-haiku-20240307"

# --- SESSION STATE ---
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "full_text" not in st.session_state: st.session_state.full_text = ""
if "analysis_report" not in st.session_state: st.session_state.analysis_report = ""
if "current_model" not in st.session_state: st.session_state.current_model = "Unknown"

# --- HELPER FUNCTIONS ---
def get_pdf_text(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def create_docx(report_text):
    """Generates a downloadable Word Document."""
    doc = Document()
    doc.add_heading('NCAAA Dentistry Accreditation Report', 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("This report analyzes the General Dentistry Program compliance with ETEC 2022 Standards.")
    
    doc.add_heading('Detailed Analysis', level=1)
    
    # Simple markdown cleanup for Word
    for line in report_text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('###') or line.startswith('**') and len(line) < 60:
            doc.add_heading(line.replace('#', '').replace('*', '').strip(), level=2)
        else:
            doc.add_paragraph(line.replace('*', ''))
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- INTELLIGENT AGENT LOGIC (Dentistry Specific) ---
def analyze_manuscript(text):
    status = st.status("🔍 Initializing Dental Accreditation Agents...", expanded=True)
    
    # 1. ATTEMPT WITH SMART MODEL (SONNET 3.5)
    try:
        status.write(f"🧠 Attempting Analysis with {SMART_MODEL}...")
        
        system_prompt = (
            "You are a Senior External Reviewer for the NCAAA (ETEC). "
            "You specialize in Bachelor of Dental Surgery (BDS) programs. "
            "You are rigorous, evidence-based, and focused on clinical safety and NQF alignment."
        )
        
        user_prompt = f"""
        EVIDENCE DOCUMENTATION:
        {text[:150000]}
        
        TASK: Write a Critical Accreditation Review Report for a Dentistry Program.
        
        **CRITICAL CHECKS:**
        1. **Clinical Safety:** Check for Radiation Safety, Infection Control, and Sharps policies.
        2. **NQF Alignment:** Verify if learning outcomes match NQF Level 6 (Bachelor).
        3. **Curriculum:** Check for balance between Pre-clinical and Clinical training.
        
        **REPORT STRUCTURE:**
        1. **Executive Decision:** (Full Accreditation / Conditional / Denial).
        2. **Standard 2 (Teaching & Learning):** Critique the clinical assessment methods.
        3. **Standard 5 (Resources):** Audit the clinic facilities and safety manuals.
        4. **Key Deficiencies:** List missing documents or evidence gaps.
        5. **Recommendations:** Specific actions to fix the gaps.
        """
        
        msg = client.messages.create(
            model=SMART_MODEL,
            max_tokens=4000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        
        st.session_state.current_model = "Claude 3.5 Sonnet (Expert Mode)"
        status.update(label="Analysis Complete (Expert Mode)", state="complete", expanded=False)
        return msg.content[0].text

    except NotFoundError:
        # 2. FALLBACK TO HAIKU (Chain of Thought)
        status.write("⚠️ Expert Model restricted. Switching to Fast Logic Mode...")
        st.session_state.current_model = "Claude 3 Haiku (Fast Mode)"
        
        # Step A: Audit
        status.write("⚙️ Step 1: Auditing Clinical Safety & NQF...")
        audit_prompt = f"Read this text: {text[:100000]}\nCheck for: 1. Infection Control Policies. 2. NQF Level 6 verbs in outcomes. Return summary."
        audit_msg = client.messages.create(
            model=FAST_MODEL, max_tokens=1000, messages=[{"role": "user", "content": audit_prompt}]
        )
        audit_summary = audit_msg.content[0].text
        
        # Step B: Final Report
        status.write("⚙️ Step 2: Writing NCAAA Report...")
        final_prompt = f"""
        EVIDENCE: {text[:100000]}
        AUDIT FINDINGS: {audit_summary}
        
        Write a Formal NCAAA Accreditation Report for a Dentistry Program.
        Focus on Standard 2 (Learning) and Standard 5 (Clinical Resources).
        """
        
        final_msg = client.messages.create(
            model=FAST_MODEL, max_tokens=4000, 
            system="You are an NCAAA Reviewer.",
            messages=[{"role": "user", "content": final_prompt}]
        )
        
        status.update(label="Analysis Complete (Fast Mode)", state="complete", expanded=False)
        return final_msg.content[0].text

    except Exception as e:
        status.update(label="Error", state="error")
        st.error(f"Unexpected Error: {e}")
        return None

# --- UI LAYOUT ---
with st.sidebar:
    st.title("🦷 NCAAA Dentistry AI")
    st.caption("General Dentistry Accreditation Platform")
    uploaded_file = st.file_uploader("Upload Evidence (PDF)", type="pdf")
    
    if st.button("Reset System"):
        st.session_state.clear()
        st.rerun()
        
    st.markdown("---")
    st.header("🖨️ Report Generator")
    if st.session_state.analysis_report:
        docx_data = create_docx(st.session_state.analysis_report)
        st.download_button(
            label="📥 Download Report (.docx)",
            data=docx_data,
            file_name="NCAAA_Dentistry_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- MAIN APP ---
if uploaded_file and not st.session_state.full_text:
    text = get_pdf_text(uploaded_file)
    if text:
        st.session_state.full_text = text
        st.success(f"Evidence Loaded: {len(text)} characters.")

if st.session_state.full_text and not st.session_state.analysis_report:
    if st.button("🚀 Start Accreditation Analysis"):
        report = analyze_manuscript(st.session_state.full_text)
        if report:
            st.session_state.analysis_report = report
            st.rerun()

if st.session_state.analysis_report:
    st.success(f"Generated using: **{st.session_state.current_model}**")
    
    tab1, tab2 = st.tabs(["📝 Compliance Report", "💬 Accreditation Consultant"])
    
    with tab1:
        st.markdown(st.session_state.analysis_report)
    
    with tab2:
        st.info("Ask specific questions about the evidence or standards.")
        
        for msg in st.session_state.chat_history:
             if msg['role'] != 'user': 
                 if len(msg['content']) < 4000:
                    st.chat_message(msg["role"]).markdown(msg["content"])
        
        if prompt := st.chat_input("Ex: 'Is the infection control policy sufficient?'"):
            st.chat_message("user").markdown(prompt)
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            
            with st.chat_message("assistant"):
                ACTIVE_MODEL = SMART_MODEL if "Sonnet" in st.session_state.current_model else FAST_MODEL
                try:
                    stream = client.messages.create(
                        model=ACTIVE_MODEL, 
                        max_tokens=2000, 
                        system="You are an Accreditation Consultant helping with the SSR.",
                        messages=[
                            {"role": "user", "content": f"Context: {st.session_state.analysis_report}"},
                            {"role": "assistant", "content": "I understand the critique."},
                            {"role": "user", "content": prompt}
                        ], 
                        stream=True
                    )
                    response = st.write_stream(chunk.delta.text for chunk in stream if chunk.type == "content_block_delta")
                except NotFoundError:
                    stream = client.messages.create(
                        model=FAST_MODEL, 
                        max_tokens=2000, 
                        system="You are an Accreditation Consultant helping with the SSR.",
                        messages=[
                            {"role": "user", "content": f"Context: {st.session_state.analysis_report}"},
                            {"role": "assistant", "content": "I understand the critique."},
                            {"role": "user", "content": prompt}
                        ], 
                        stream=True
                    )
                    response = st.write_stream(chunk.delta.text for chunk in stream if chunk.type == "content_block_delta")

            st.session_state.chat_history.append({"role": "assistant", "content": response})
else:
    if not uploaded_file: st.info("👈 Upload PDF to begin.")
